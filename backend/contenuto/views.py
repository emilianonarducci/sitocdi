import json
import os
import io
import threading
from django.http import JsonResponse
from django.shortcuts import render, redirect
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib import messages
from django.db import connection
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from .models import HeroSlide, ConsigliatiCard, PercheSceglierciSezione, SalutePerTeArticolo, SchedaMedico, CVImportJob, NewsletterIscritto


def hero_api(request):
    slides = HeroSlide.objects.filter(attivo=True)
    return JsonResponse([{
        "id": s.id,
        "badge": s.badge,
        "titolo": s.titolo,
        "sottotitolo": s.sottotitolo,
        "cta_testo": s.cta_testo,
        "cta_link": s.cta_link,
        "immagine_url": s.immagine_url,
    } for s in slides], safe=False)


def consigliati_api(request):
    cards = ConsigliatiCard.objects.filter(attivo=True)
    return JsonResponse([{
        "id": c.id,
        "titolo": c.titolo,
        "immagine_url": c.immagine_url,
        "link": c.link,
    } for c in cards], safe=False)


def perche_sceglierci_api(request):
    try:
        sezione = PercheSceglierciSezione.objects.first()
        if not sezione:
            return JsonResponse({}, safe=False)
        return JsonResponse({
            "titolo": sezione.titolo,
            "descrizione": sezione.descrizione,
            "immagine_sfondo_url": sezione.immagine_sfondo_url,
            "cards": [{
                "id": c.id,
                "variante": c.variante,
                "titolo": c.titolo,
                "descrizione": c.descrizione,
            } for c in sezione.cards.all()],
        })
    except Exception:
        return JsonResponse({}, safe=False)


def salute_per_te_api(request):
    tab = request.GET.get("tab")
    qs = SalutePerTeArticolo.objects.filter(attivo=True)
    if tab:
        qs = qs.filter(tab=tab)
    return JsonResponse([{
        "id": a.id,
        "tab": a.tab,
        "titolo": a.titolo,
        "descrizione": a.descrizione,
        "immagine_url": a.immagine_url,
        "link": a.link,
    } for a in qs], safe=False)


def _medico_to_dict(m: SchedaMedico) -> dict:
    return {
        "id": m.id,
        "nome_completo": m.nome_completo,
        "specializzazione": m.specializzazione,
        "ruolo": m.ruolo,
        "foto_url": m.foto_url,
        "bio": m.bio,
        "titoli_accademici": m.titoli_accademici,
        "esperienze_professionali": m.esperienze_professionali,
        "pubblicazioni": m.pubblicazioni,
        "link_prenota": m.link_prenota,
    }


@csrf_exempt
@require_POST
def newsletter_subscribe(request):
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({"error": "JSON non valido"}, status=400)

    email = (data.get("email") or "").strip().lower()
    if not email:
        return JsonResponse({"error": "Email obbligatoria"}, status=400)

    nome = (data.get("nome") or "").strip()
    cognome = (data.get("cognome") or "").strip()

    obj, created = NewsletterIscritto.objects.get_or_create(
        email=email,
        defaults={"nome": nome, "cognome": cognome, "attivo": True},
    )
    if not created and not obj.attivo:
        obj.attivo = True
        obj.nome = nome or obj.nome
        obj.cognome = cognome or obj.cognome
        obj.save(update_fields=["attivo", "nome", "cognome"])

    return JsonResponse({"ok": True, "created": created})


def medici_list_api(request):
    qs = SchedaMedico.objects.filter(attivo=True)
    return JsonResponse([_medico_to_dict(m) for m in qs], safe=False)


def medico_detail_api(request, pk: int):
    try:
        m = SchedaMedico.objects.get(pk=pk, attivo=True)
        return JsonResponse(_medico_to_dict(m))
    except SchedaMedico.DoesNotExist:
        return JsonResponse({"error": "Non trovato"}, status=404)


def _extract_text(file_bytes: bytes, filename: str) -> str:
    if filename.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif filename.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        # Europass CVs store most content in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt and txt not in parts:
                        parts.append(txt)
        return "\n".join(parts)
    elif filename.endswith(".doc"):
        import subprocess, tempfile
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        result = subprocess.run(["antiword", tmp_path], capture_output=True, text=True)
        os.unlink(tmp_path)
        if result.returncode != 0:
            raise Exception(result.stderr or "antiword ha restituito un errore")
        return result.stdout
    raise Exception("Formato non supportato. Usa PDF, DOC o DOCX.")


def _process_cv_job(job_id: int, file_bytes: bytes, filename: str):
    """Background thread: extract text, call Claude, save SchedaMedico, update job."""
    try:
        job = CVImportJob.objects.get(pk=job_id)
        job.status = CVImportJob.STATUS_PROCESSING
        job.save(update_fields=["status"])

        # 1. Extract text
        text = _extract_text(file_bytes, filename)
        if not text.strip():
            raise Exception("Il file non contiene testo estraibile (potrebbe essere una scansione).")

        # 2. Call Claude
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            raise Exception("ANTHROPIC_API_KEY non configurata nel server.")

        system_prompt = """Sei un assistente specializzato nell'estrazione di dati strutturati da curriculum vitae di medici italiani.
Il CV potrebbe essere in formato Europass, formato libero, o con contenuto distribuito in tabelle.
Estrai tutte le informazioni disponibili con la massima precisione.
Rispondi SEMPRE e SOLO con un oggetto JSON valido, senza markdown, senza commenti, senza testo aggiuntivo."""

        prompt = f"""Analizza il seguente CV medico ed estrai i dati nel formato JSON specificato.

REGOLE IMPORTANTI:
- Tutti i campi stringa devono essere stringhe non null; usa "" se l'informazione non è presente
- "titoli_accademici" include: lauree, specializzazioni, master, dottorati, corsi di perfezionamento
- "esperienze_professionali" include: posizioni lavorative, incarichi ospedalieri, attività ambulatoriali
- "bio" deve essere una presentazione professionale di 3-5 frasi scorrevoli in italiano, in terza persona
- "specializzazione" è la branca medica principale (es. "Cardiologia", "Ortopedia e Traumatologia")
- "ruolo" è la qualifica attuale (es. "Cardiologo Strutturato", "Primario", "Medico Specialista")
- "pubblicazioni" è un testo con le pubblicazioni principali separate da newline; usa "" se assenti
- Negli array, ordina dal più recente al meno recente

FORMATO JSON RICHIESTO:
{{
  "nome_completo": "Titolo Nome Cognome",
  "specializzazione": "Specializzazione medica principale",
  "ruolo": "Ruolo / qualifica attuale",
  "bio": "Presentazione professionale in terza persona...",
  "titoli_accademici": [
    {{"anno": "2014 – 2018", "titolo": "Specializzazione in Cardiologia (110/110 con lode) – Università degli Studi di Milano"}}
  ],
  "esperienze_professionali": [
    {{"anno": "2018 – presente", "descrizione": "Cardiologo Strutturato – ASST Ospedale Niguarda, Milano. Responsabile ambulatorio scompenso cardiaco."}}
  ],
  "pubblicazioni": "Ferretti MA et al. (2022). Long-term outcomes... European Heart Journal.\nFerretti MA et al. (2020)..."
}}

CV DA ANALIZZARE:
---
{text[:14000]}
---"""

        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=3000,
            system=system_prompt,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = response.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        data = json.loads(raw)

        # 3. Save SchedaMedico
        scheda = SchedaMedico.objects.create(
            nome_completo=data.get("nome_completo") or "",
            specializzazione=data.get("specializzazione") or "",
            ruolo=data.get("ruolo") or "",
            bio=data.get("bio") or "",
            titoli_accademici=data.get("titoli_accademici") or [],
            esperienze_professionali=data.get("esperienze_professionali") or [],
            pubblicazioni=data.get("pubblicazioni") or "",
            attivo=True,
        )

        job.scheda = scheda
        job.status = CVImportJob.STATUS_DONE
        job.save(update_fields=["scheda", "status"])

    except Exception as e:
        CVImportJob.objects.filter(pk=job_id).update(
            status=CVImportJob.STATUS_ERROR,
            error_message=str(e),
        )
    finally:
        connection.close()


@staff_member_required
def upload_cv_pdf(request):
    if request.method == "GET":
        return render(request, "contenuto/upload_cv.html")

    pdf_file = request.FILES.get("pdf")
    if not pdf_file:
        messages.error(request, "Nessun file selezionato.")
        return render(request, "contenuto/upload_cv.html")

    filename = pdf_file.name.lower()
    if not any(filename.endswith(ext) for ext in (".pdf", ".doc", ".docx")):
        messages.error(request, "Formato non supportato. Usa PDF, DOC o DOCX.")
        return render(request, "contenuto/upload_cv.html")

    file_bytes = pdf_file.read()
    job = CVImportJob.objects.create(filename=pdf_file.name)

    t = threading.Thread(target=_process_cv_job, args=(job.id, file_bytes, filename), daemon=True)
    t.start()

    return redirect(f"/api/cms/upload-cv/status/{job.id}/")


@staff_member_required
def cv_import_status(request, job_id: int):
    try:
        job = CVImportJob.objects.get(pk=job_id)
    except CVImportJob.DoesNotExist:
        return JsonResponse({"error": "Job non trovato"}, status=404)

    if request.headers.get("Accept") == "application/json":
        return JsonResponse({
            "status": job.status,
            "scheda_id": job.scheda_id,
            "error_message": job.error_message,
        })

    return render(request, "contenuto/cv_import_status.html", {"job": job})


@staff_member_required
def gdrive_newsletter_sync(request):
    if request.method == "GET":
        return render(request, "contenuto/gdrive_newsletter_sync.html")

    # POST → esegui sync
    folder_id = os.environ.get("GDRIVE_FOLDER_ID", "")
    has_creds = os.environ.get("GDRIVE_SERVICE_ACCOUNT_JSON") or os.environ.get("GDRIVE_SERVICE_ACCOUNT_FILE")

    if not has_creds or not folder_id:
        messages.error(request, "Variabili d'ambiente mancanti: GDRIVE_FOLDER_ID e credenziali service account")
        return render(request, "contenuto/gdrive_newsletter_sync.html")

    try:
        from .management.commands.sync_newsletter_gdrive import run_sync
        result = run_sync(folder_id)

        if not result["files"]:
            messages.warning(request, result.get("message", "Nessun file trovato"))
        else:
            file_names = ", ".join(f["name"] for f in result["files"])
            messages.success(
                request,
                f"Sync completata — {len(result['files'])} file: {file_names} | "
                f"Nuovi: {result['created']} · Aggiornati: {result['updated']} · Saltati: {result['skipped']}"
            )
        if result["errors"]:
            for err in result["errors"]:
                messages.warning(request, f"Riga ignorata: {err}")
    except Exception as e:
        messages.error(request, f"Errore durante la sync: {e}")

    return render(request, "contenuto/gdrive_newsletter_sync.html")
